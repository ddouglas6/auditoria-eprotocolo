import streamlit as st
import pandas as pd
import pdfplumber
import io
import re
import time
import json
import os
import requests
import urllib3
import base64
from datetime import datetime
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select
from fpdf import FPDF
import docx

# Desativa avisos de SSL em sites do Governo
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# ==========================================
# 1. CONFIGURAÇÃO DA PÁGINA E ESTILO SAAS
# ==========================================
st.set_page_config(page_title="Auditoria e-Protocolo", page_icon="🛡️", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif !important; }
    [data-testid="stConnectionStatus"] {display: none !important;}
    .main-header {background: linear-gradient(135deg, #2563eb 0%, #1e40af 100%); padding: 30px; border-radius: 15px; color: white; text-align: center; margin-bottom: 25px; box-shadow: 0 10px 30px -10px rgba(37, 99, 235, 0.4);}
    .main-header h1 {margin: 0; font-size: 2.2em; font-weight: 700; color: white; letter-spacing: -0.5px;}
    .log-box {background-color: #0f172a; color: #10b981; padding: 20px; border-radius: 12px; font-family: 'SFMono-Regular', Consolas, monospace; font-size: 13px; height: 300px; overflow-y: auto; box-shadow: inset 0 2px 4px rgba(0,0,0,0.5);}
    .stButton>button {border-radius: 12px; font-weight: 600; padding: 0.6rem 1rem;}
    .stButton>button[kind="primary"] {background-color: #2563eb; color: white; font-size: 1.1em;}
    </style>
    <script>
    async function keepAwake() { if ('wakeLock' in navigator) { try { await navigator.wakeLock.request('screen'); } catch (err) {} } }
    keepAwake(); document.addEventListener('visibilitychange', () => { if (document.visibilityState === 'visible') { keepAwake(); } });
    </script>
""", unsafe_allow_html=True)

# ==========================================
# 2. INICIALIZAÇÃO BLINDADA DO ESTADO
# ==========================================
PARAMETROS_PADRAO = {
    "usr": "", "pwd": "", "gemini_key": "", "filtro_dias": 0, "alerta_dias": 30,
    "cb_somente_nao_atribuidos": False, "cb_historico": False,
    "cb_pdf": True, "cb_excel": True, "cb_word": False, "ordem_relatorio": "Decrescente",
    "cb_ia_risco": True, "cb_resumo_ia": False, "txt_palavra_ia": "Mandado", "dd_prioridade_ia": "🔴 ALTO"
}

if "app_fase" not in st.session_state:
    st.session_state.app_fase = "configuracao"
    st.session_state.resultados = []
    st.session_state.erro_msg = ""
    st.session_state.arquivo_importado_id = None
    for k, v in PARAMETROS_PADRAO.items():
        st.session_state[k] = v

# ==========================================
# 3. GATILHO DE EVENTO (O Segredo do Clique)
# ==========================================
def acionar_auditoria_callback():
    """Garante que a página mude de estado antes de recarregar a tela (Evita cliques fantasmas)."""
    u = str(st.session_state.usr).strip()
    p = str(st.session_state.pwd).strip()
    
    if not u or not p:
        st.session_state.erro_msg = "⚠️ Atenção: Preencha seu CPF e Senha."
    else:
        st.session_state.erro_msg = ""
        st.session_state.app_fase = "processando"

def reiniciar_app_callback():
    st.session_state.app_fase = "configuracao"
    st.session_state.resultados = []
    st.session_state.erro_msg = ""

# ==========================================
# 4. FUNÇÕES DE NEGÓCIO E IA
# ==========================================
def processar_texto_limpo(texto):
    if not isinstance(texto, str): texto = str(texto)
    return texto.replace('“', '"').replace('”', '"').replace('‘', "'").replace('’', "'").replace('–', '-').replace('—', '-').encode('latin-1', 'replace').decode('latin-1')

def calcular_dias_parado(data_str):
    try: 
        match = re.search(r'\d{2}/\d{2}/\d{4}', str(data_str))
        if match: return max(0, (datetime.now() - datetime.strptime(match.group(), "%d/%m/%Y")).days)
    except: pass
    return 0

def consultar_gemini(texto_pdf, api_key):
    if not api_key or len(str(texto_pdf).strip()) < 10: return "Sem IA ou sem texto."
    try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
        prompt = f"Resuma o documento a seguir em 1 parágrafo contendo: Assunto, Próximo Passo, Último Andamento.\n\n{str(texto_pdf)[:30000]}"
        payload = {"contents": [{"parts": [{"text": prompt}]}]}
        resp = requests.post(url, json=payload, timeout=20)
        if resp.status_code == 200: return resp.json()['candidates'][0]['content']['parts'][0]['text'].strip().replace('*', '')
        return f"Falha na API (Erro {resp.status_code})"
    except: return "Falha de Conexão com IA."

def analisar_risco_simples(texto, palavra_chave, prioridade):
    texto_baixo = str(texto).lower()
    if palavra_chave.strip() and palavra_chave.lower().strip() in texto_baixo: return prioridade
    if any(p in texto_baixo for p in ['urgente', 'imediato', 'mandado', 'judicial', 'liminar']): return "🔴 ALTO"
    if any(p in texto_baixo for p in ['solicitação', 'memorando', 'informação']): return "🟡 MÉDIO"
    return "🟢 BAIXO"

# ==========================================
# 5. MOTOR SELENIUM DE EXTRAÇÃO
# ==========================================
def iniciar_auditoria_backend():
    download_dir = "/tmp/downloads_eprotocolo"
    os.makedirs(download_dir, exist_ok=True)
    
    options = Options()
    options.add_argument('--headless=new')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--window-size=1920,1080')
    options.add_argument('--disable-blink-features=AutomationControlled')
    options.add_experimental_option("excludeSwitches", ["enable-automation"])
    options.add_experimental_option('useAutomationExtension', False)
    options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")
    options.add_experimental_option("prefs", {"download.default_directory": download_dir, "download.prompt_for_download": False, "plugins.always_open_pdf_externally": True, "pdfjs.disabled": True})
    
    driver = webdriver.Chrome(options=options)
    driver.execute_cdp_cmd('Page.setDownloadBehavior', {'behavior': 'allow', 'downloadPath': download_dir})
    driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
    wait = WebDriverWait(driver, 45)
    dados_coletados = []
    
    try:
        yield "🚪 Acessando portal de Autenticação...", 10
        driver.get("https://auth-cs.identidadedigital.pr.gov.br/centralautenticacao/login.html?response_type=code&client_id=9188905e74c28e489b44e954ec0b9bca&redirect_uri=https%3A%2F%2Fwww.eprotocolo.pr.gov.br%2Fspiweb")
        time.sleep(3)
        
        try: driver.find_element(By.ID, "btnCentral").click(); time.sleep(2)
        except: pass
        
        yield "🔑 Injetando credenciais...", 15
        driver.find_element(By.ID, "attribute_central").send_keys(str(st.session_state.usr).strip())
        time.sleep(1)
        driver.find_element(By.ID, "password").send_keys(str(st.session_state.pwd).strip())
        time.sleep(1)
        driver.find_element(By.ID, "btn-central-acessar").click()
        
        yield "⏳ Aguardando liberação do Servidor...", 25
        for _ in range(30):
            if "telaInicial" in driver.current_url or "iniciarProcesso" in driver.current_url: break
            time.sleep(1)
            
        if "login" in driver.current_url.lower() or "centralautenticacao" in driver.current_url.lower():
            raise Exception("Falha de Autenticação. Verifique seu CPF/Senha ou se o site do Estado está em manutenção.")
            
        try: driver.get("https://www.eprotocolo.pr.gov.br/spiweb/telaInicial.do?action=iniciarProcesso")
        except: driver.execute_script("window.stop();")
        time.sleep(4)
        
        yield "📍 Acessando 'Protocolos no Local'...", 35
        driver.find_element(By.XPATH, "//a[@href='#div_protocolos' or contains(text(), 'Protocolos no Local')]").click()
        time.sleep(5)
        
        select_locais = Select(wait.until(EC.presence_of_element_located((By.ID, "codLocal"))))
        lista_locais = [opt.text.strip() for opt in select_locais.options if "Selecione" not in opt.text]
        janela_principal = driver.current_window_handle
        
        yield f"📋 {len(lista_locais)} Setores Encontrados. Iniciando varredura...", 40
        progresso_por_local = 50 / max(len(lista_locais), 1)
        progresso_atual = 40
        
        for nome_local in lista_locais:
            progresso_atual += progresso_por_local
            yield f"🔎 Buscando no Setor: {nome_local}", progresso_atual
            
            Select(wait.until(EC.presence_of_element_located((By.ID, "codLocal")))).select_by_visible_text(nome_local)
            time.sleep(2)
            try: driver.find_element(By.ID, "botaoPesquisar").click()
            except: pass
            time.sleep(5)
            
            linhas_tabela = driver.find_elements(By.XPATH, "//table//tbody//tr[.//img[contains(@src, 'icon_exibir.svg')]]")
            if not linhas_tabela: continue
            
            for idx in range(len(linhas_tabela)):
                try:
                    linhas_tabela = driver.find_elements(By.XPATH, "//table//tbody//tr[.//img[contains(@src, 'icon_exibir.svg')]]")
                    linha = linhas_tabela[idx]
                    
                    try: atribuido_para = linha.find_element(By.XPATH, ".//*[contains(@id, '_atribuidoPara')]").text.strip() or "-"
                    except: atribuido_para = "-"
                    
                    if st.session_state.cb_somente_nao_atribuidos and (atribuido_para != "-"): continue
                        
                    driver.execute_script("arguments[0].click();", linha.find_element(By.XPATH, ".//img[contains(@src, 'icon_exibir.svg')]/ancestor::a"))
                    time.sleep(4)
                    
                    def obter_texto(xpath):
                        try: return driver.find_element(By.XPATH, xpath).text.strip()
                        except: return "-"
                        
                    try: protocolo_id = driver.find_element(By.ID, "numeroProtocolo").get_attribute("value")
                    except: protocolo_id = "-"
                    
                    data_envio = obter_texto("//div[contains(text(), 'Enviado em:')]/following-sibling::div[1]")
                    dias_parado = calcular_dias_parado(data_envio)
                    
                    if dias_parado >= st.session_state.filtro_dias:
                        yield f"      📄 ID {protocolo_id} capturado ({dias_parado} dias parado).", progresso_atual
                        
                        detalhamento = obter_texto("//div[contains(text(), 'Detalhamento:')]/following-sibling::div[1]")
                        dados = {
                            "Setor": nome_local, "Protocolo": protocolo_id, "Atribuído Para": atribuido_para,
                            "Enviado Em": data_envio, "Dias Parado": dias_parado,
                            "Situação": obter_texto("//div[contains(text(), 'Situação:')]/following-sibling::div[1]"),
                            "Local de Envio": obter_texto("//div[contains(text(), 'Local de Envio:')]/following-sibling::div[1]"),
                            "Motivo": obter_texto("//div[contains(text(), 'Motivo:')]/following-sibling::div[1]"),
                            "Detalhamento": detalhamento
                        }
                        
                        if st.session_state.cb_ia_risco:
                            dados["Risco Calculado"] = analisar_risco_simples(detalhamento, st.session_state.txt_palavra_ia, st.session_state.dd_prioridade_ia)
                            
                        if st.session_state.cb_resumo_ia:
                            yield f"      🤖 Extraindo Anexo para IA...", progresso_atual
                            try:
                                for f in os.listdir(download_dir): os.remove(os.path.join(download_dir, f))
                                driver.execute_script("arguments[0].click();", wait.until(EC.presence_of_element_located((By.XPATH, "//a[.//img[contains(@src, 'icon_download.svg')]] | //img[contains(@src, 'icon_download.svg')]"))))
                                
                                path_pdf = None
                                for _ in range(45):
                                    time.sleep(1)
                                    arqs = [f for f in os.listdir(download_dir) if not f.endswith('.crdownload') and not f.endswith('.tmp')]
                                    if arqs:
                                        arqs.sort(key=lambda x: os.path.getmtime(os.path.join(download_dir, x)), reverse=True)
                                        path_pdf = os.path.join(download_dir, arqs[0]); time.sleep(1); break
                                        
                                if path_pdf:
                                    with pdfplumber.open(path_pdf) as pdf:
                                        paginas = pdf.pages if len(pdf.pages) <= 20 else pdf.pages[:10] + pdf.pages[-10:]
                                        texto_documento = "\n".join([page.extract_text() or "" for page in paginas])
                                    dados["Resumo IA (Anexo)"] = consultar_gemini(texto_documento, st.session_state.gemini_key)
                                else: dados["Resumo IA (Anexo)"] = "Falha no Download."
                            except: dados["Resumo IA (Anexo)"] = "PDF Bloqueado/Inexistente."
                            finally:
                                while len(driver.window_handles) > 1:
                                    driver.switch_to.window(driver.window_handles[-1])
                                    driver.close()
                                driver.switch_to.window(janela_principal)

                        if st.session_state.cb_historico:
                            try:
                                driver.find_element(By.XPATH, "//h2[contains(., 'Andamentos')]").click(); time.sleep(2)
                                dados["Movimentações"] = len(driver.find_elements(By.XPATH, "//div[@id='Andamentos_menos']//table//tbody//tr"))
                            except: dados["Movimentações"] = 0
                            
                        dados_coletados.append(dados)
                        
                except Exception: pass
                finally:
                    try:
                        if driver.current_window_handle != janela_principal: driver.switch_to.window(janela_principal)
                        btn_voltar = wait.until(EC.element_to_be_clickable((By.XPATH, "//input[@value='Voltar'] | //button[contains(text(), 'Voltar')]")))
                        driver.execute_script("arguments[0].click();", btn_voltar)
                    except: driver.execute_script("window.history.go(-1)")
                    time.sleep(2)
        
        yield "✅ Processamento Concluído com Sucesso!", 100
        return dados_coletados
        
    except Exception as erro: raise erro
    finally:
        try: driver.quit()
        except: pass

# ==========================================
# 6. ESTRUTURA VISUAL DA APLICAÇÃO (TELAS)
# ==========================================

st.markdown('<div class="main-header"><h1>🛡️ Auditoria de e-Protocolo</h1><p>Motor de Produção • Limpo e Autônomo</p></div>', unsafe_allow_html=True)

# ----------------- SIDEBAR -----------------
with st.sidebar:
    st.markdown("### 💾 Gestão de Perfil")
    arquivo_importado = st.file_uploader("Importar Perfil (.json)", type="json", label_visibility="collapsed")
    
    # Processa o arquivo silenciosamente, convertendo CPF e Senha para TEXTO (string)
    if arquivo_importado and st.session_state.arquivo_importado_id != arquivo_importado.file_id:
        try:
            dados_lidos = json.loads(arquivo_importado.getvalue().decode('utf-8'))
            for chave, valor in dados_lidos.items():
                if chave in PARAMETROS_PADRAO:
                    # Trava de Segurança: Força as credenciais e IA a serem texto
                    if chave in ["usr", "pwd", "gemini_key", "txt_palavra_ia"]:
                        st.session_state[chave] = str(valor)
                    else:
                        st.session_state[chave] = valor
                        
            st.session_state.arquivo_importado_id = arquivo_importado.file_id
            st.success("Perfil lido com sucesso!")
        except Exception: pass

    # Exportar perfil
    dados_exportacao = {k: st.session_state[k] for k in PARAMETROS_PADRAO.keys()}
    st.download_button("💾 Exportar Configurações", json.dumps(dados_exportacao, indent=4), "perfil_eprotocolo.json", "application/json", use_container_width=True)
    
    st.divider()
    st.markdown("### 🔐 Credenciais")
    st.text_input("CPF do Servidor:", key="usr")
    st.text_input("Senha Governamental:", key="pwd", type="password")
    
    st.divider()
    st.markdown("### 🤖 Cérebro Artificial")
    st.text_input("Chave API do Gemini:", key="gemini_key", type="password")
    st.checkbox("Analisar Risco Base", key="cb_ia_risco")
    st.checkbox("Ler e Resumir PDFs", key="cb_resumo_ia")
    with st.expander("Parâmetros do Risco"):
        st.text_input("Palavra-Chave Crítica:", key="txt_palavra_ia")
        st.selectbox("Nível de Alerta:", ["🔴 ALTO", "🟡 MÉDIO", "🟢 BAIXO"], key="dd_prioridade_ia")

    # BOTÃO PARA CELULAR: Fica no menu para evitar "clique fantasma"
    st.divider()
    st.button("🚀 INICIAR VARREDURA (Mobile)", type="primary", use_container_width=True, on_click=acionar_auditoria_callback, key="btn_mobile")

# ----------------- FLUXO DE TELAS PRINCIPAIS -----------------

if st.session_state.app_fase == "configuracao":
    
    if st.session_state.erro_msg:
        st.error(st.session_state.erro_msg)
        
    st.markdown("### ⚙️ Ajustes da Varredura")
    c1, c2 = st.columns(2)
    with c1:
        st.number_input("Ignorar dias mais recentes:", min_value=0, key="filtro_dias")
        st.number_input("Destacar atraso (em dias):", min_value=1, key="alerta_dias")
        st.checkbox("Somente Protocolos Não Atribuídos", key="cb_somente_nao_atribuidos")
        st.checkbox("Coletar Histórico (Qtd Andamentos)", key="cb_historico")
    with c2:
        st.checkbox("Relatório PDF", key="cb_pdf")
        st.checkbox("Planilha Excel", key="cb_excel")
        st.checkbox("Ofício Word", key="cb_word")
        st.radio("Ordenação:", ["Decrescente", "Crescente"], key="ordem_relatorio")

    st.markdown("<br>", unsafe_allow_html=True)
    # BOTÃO PRINCIPAL COM CALLBACK DIRETO
    st.button("🚀 INICIAR VARREDURA", type="primary", use_container_width=True, on_click=acionar_auditoria_callback, key="btn_pc")

elif st.session_state.app_fase == "processando":
    st.info("⚠️ O Motor Selenium está rodando em plano de fundo. Deixe a tela ligada.")
    barra_progresso = st.progress(0)
    caixa_texto_logs = st.empty()
    historico_logs = []
    
    try:
        motor = iniciar_auditoria_backend()
        while True:
            try:
                mensagem, progresso_num = next(motor)
                historico_logs.append(f"[{datetime.now().strftime('%H:%M:%S')}] {mensagem}")
                caixa_texto_logs.markdown(f"<div class='log-box'>{'<br>'.join(historico_logs[-15:])}</div>", unsafe_allow_html=True)
                barra_progresso.progress(int(progresso_num) / 100.0)
            except StopIteration as retorno:
                st.session_state.resultados = retorno.value
                st.session_state.app_fase = "sucesso"
                st.rerun()
                break
    except Exception as e:
        st.session_state.erro_msg = str(e)
        st.session_state.app_fase = "erro"
        st.rerun()

elif st.session_state.app_fase == "sucesso":
    st.success("🎉 Processo concluído sem interrupções.")
    
    df = pd.DataFrame(st.session_state.resultados)
    if len(df) == 0:
        st.warning("A busca ocorreu perfeitamente, mas nenhum protocolo atendeu aos seus filtros.")
    else:
        crescente = (st.session_state.ordem_relatorio == "Crescente")
        df = df.sort_values(by=['Setor', 'Dias Parado'], ascending=[True, crescente])
        
        with st.expander("👁️ Ver Dados Brutos", expanded=True):
            st.dataframe(df, use_container_width=True)
            
        st.markdown("### 📥 Documentos Gerados")
        btn_col1, btn_col2, btn_col3 = st.columns(3)
        
        if st.session_state.cb_excel:
            buffer_excel = io.BytesIO()
            df.to_excel(buffer_excel, index=False)
            btn_col1.download_button("📊 Baixar Tabela Excel", buffer_excel.getvalue(), "Auditoria_Resultados.xlsx", "application/vnd.ms-excel", use_container_width=True)

        if st.session_state.cb_pdf:
            pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial", 'B', 16); pdf.cell(0, 10, txt="Relatório Oficial - Auditoria", ln=True, align='C')
            agrupamento = {}
            for _, row in df.iterrows(): agrupamento.setdefault(row['Setor'], []).append(row)
            
            for setor, protocolos in agrupamento.items():
                pdf.ln(5); pdf.set_font("Arial", 'B', 12); pdf.set_fill_color(37, 99, 235); pdf.set_text_color(255, 255, 255)
                pdf.multi_cell(0, 8, processar_texto_limpo(f" SETOR: {setor} ({len(protocolos)} Identificados)"), fill=True)
                for prot in protocolos:
                    atrasado = prot.get('Dias Parado', 0) >= st.session_state.alerta_dias
                    pdf.set_font("Arial", 'B', 10); pdf.set_text_color(220, 38, 38) if atrasado else pdf.set_text_color(0, 0, 0)
                    pdf.ln(3); pdf.cell(0, 6, processar_texto_limpo(f"Protocolo: {prot.get('Protocolo')} | Dias sem Movimento: {prot.get('Dias Parado')}"), 0, 1)
                    pdf.set_font("Arial", '', 10); pdf.set_text_color(30, 41, 59)
                    for k, v in prot.items():
                        if k not in ["Protocolo", "Setor", "Dias Parado"]: pdf.multi_cell(0, 5, processar_texto_limpo(f"{k}: {v}"))
                    pdf.set_draw_color(203, 213, 225); pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            btn_col2.download_button("📕 Baixar Livro PDF", pdf.output(dest='S').encode('latin-1'), "Auditoria_Resultados.pdf", "application/pdf", use_container_width=True)

        if st.session_state.cb_word:
            doc = docx.Document(); doc.add_heading('Relatório Oficial de Auditoria', 0).alignment = 1 
            for _, prot in df.iterrows():
                doc.add_heading(f"Protocolo: {prot.get('Protocolo')}", level=2)
                for k, v in prot.items():
                    if k != "Protocolo": doc.add_paragraph(f"{k}: {v}")
            buffer_word = io.BytesIO(); doc.save(buffer_word)
            btn_col3.download_button("📘 Baixar Documento Word", buffer_word.getvalue(), "Auditoria_Resultados.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

    st.markdown("<br>", unsafe_allow_html=True)
    st.button("🔄 Retornar e Fazer Nova Varredura", type="secondary", use_container_width=True, on_click=reiniciar_app_callback)

elif st.session_state.app_fase == "erro":
    st.error("❌ Ocorreu uma interrupção inesperada no motor de captura.")
    st.warning("Diagnóstico Técnico:")
    st.code(st.session_state.erro_msg, language='text')
    st.button("🔄 Voltar ao Início", type="primary", on_click=reiniciar_app_callback)
